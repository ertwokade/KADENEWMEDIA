from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "public" / "downloads" / "kadeai-sirket-briefi-ornek.docx"

INK = RGBColor(24, 24, 27)
BODY = RGBColor(39, 39, 42)
MUTED = RGBColor(82, 82, 91)
AMBER = RGBColor(154, 114, 0)
AMBER_BRIGHT = "F2C322"
AMBER_PALE = "FFF8D6"
LINE = "E4E4E7"


def set_run_font(run, size=None, color=BODY, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_paragraph_border(paragraph, side, color, size=10, space=6):
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_run_font(run, 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_label_value(document, label, value):
    paragraph = document.add_paragraph(style="Kade Label Detail")
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, 10.5, INK, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, 10.5, BODY)
    return paragraph


def create_bullet_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    number_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.extend([start, number_format, level_text, suffix, justification])

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.extend([tabs, indent])
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return number_id


def add_bullets(document, items, number_id):
    for item in items:
        paragraph = document.add_paragraph(style="Kade Bullet")
        properties = paragraph._p.get_or_add_pPr()
        number_properties = OxmlElement("w:numPr")
        level = OxmlElement("w:ilvl")
        level.set(qn("w:val"), "0")
        number = OxmlElement("w:numId")
        number.set(qn("w:val"), str(number_id))
        number_properties.extend([level, number])
        properties.append(number_properties)
        run = paragraph.add_run(item)
        set_run_font(run, 10.5, BODY)


def add_section(document, title, page_break_before=False):
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.page_break_before = page_break_before
    run = paragraph.add_run(title)
    set_run_font(run, 16, AMBER, bold=True)
    return paragraph


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = AMBER
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = INK
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.keep_with_next = True

    kade_bullet = styles.add_style("Kade Bullet", WD_STYLE_TYPE.PARAGRAPH)
    kade_bullet.base_style = normal
    kade_bullet.paragraph_format.space_after = Pt(4)
    kade_bullet.paragraph_format.line_spacing = 1.25

    detail = styles.add_style("Kade Label Detail", WD_STYLE_TYPE.PARAGRAPH)
    detail.base_style = normal
    detail.paragraph_format.space_after = Pt(5)
    detail.paragraph_format.line_spacing = 1.25
    detail.paragraph_format.keep_together = True


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_styles(document)
    bullet_number_id = create_bullet_numbering(document)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header.add_run("KADEAI  /  ŞİRKET HAFIZASI")
    set_run_font(left, 8.5, AMBER, bold=True)
    right = header.add_run("\tÖRNEK BRIEF")
    set_run_font(right, 8.5, MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    add_page_number(footer)
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.paragraph_format.space_before = Pt(0)
    add_page_number(first_footer)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(30)
    kicker.paragraph_format.space_after = Pt(2)
    kicker_run = kicker.add_run("KadeAI şirket hafızası şablonu")
    set_run_font(kicker_run, 10, AMBER, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run("Şirket Briefi Örneği")
    set_run_font(title_run, 29, INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        "KadeAI'nin markanızı, hedef kitlenizi ve içerik kurallarınızı tanıması için düzenlenebilir örnek."
    )
    set_run_font(subtitle_run, 12.5, MUTED)

    note = document.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.12)
    note.paragraph_format.right_indent = Inches(0.12)
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(14)
    note.paragraph_format.line_spacing = 1.2
    shade_paragraph(note, AMBER_PALE)
    set_paragraph_border(note, "left", AMBER_BRIGHT, size=18, space=8)
    note_run = note.add_run(
        "Kullanım: Bu örnekteki bilgileri kendi şirketinizle değiştirin. Başlıkları koruyun ve dosyayı KadeAI > Ayarlar > Şirket Briefi alanından yükleyin."
    )
    set_run_font(note_run, 10.5, INK, bold=True)

    add_section(document, "Marka özeti")
    add_label_value(document, "Marka adı", "Lunera Coffee")
    add_label_value(document, "Web sitesi", "https://www.luneracoffee.example")
    add_label_value(
        document,
        "Sektör / niş",
        "Yeni nesil kahve, sürdürülebilir çekirdek tedariği ve ev baristalığı",
    )

    add_section(document, "Hedef kitle")
    audience = document.add_paragraph()
    audience_run = audience.add_run(
        "22-40 yaş arası, şehirde yaşayan, nitelikli kahveye ilgi duyan; evde daha iyi kahve hazırlamak isteyen ve sürdürülebilir markaları tercih eden tüketiciler."
    )
    set_run_font(audience_run, 10.5, BODY)

    add_section(document, "Marka tonu")
    voice = document.add_paragraph()
    voice_run = voice.add_run(
        "Bilgili ama ukala olmayan, sıcak, merak uyandıran ve gündelik Türkçe kullanan. Teknik bilgileri sade örneklerle açıklar."
    )
    set_run_font(voice_run, 10.5, BODY)

    add_section(document, "Ürün ve hizmetler")
    add_bullets(
        document,
        [
            "Tek kökenli kahve çekirdekleri",
            "Aylık kahve aboneliği",
            "Ev baristalığı atölyeleri",
            "Demleme ekipmanları",
        ],
        bullet_number_id,
    )

    add_section(document, "İçerik hedefleri", page_break_before=True)
    add_bullets(
        document,
        [
            "Kahve eğitiminde güvenilir başvuru noktası olmak",
            "Abonelik deneme dönüşümlerini artırmak",
            "Atölye kayıtlarını büyütmek",
            "Topluluk etkileşimini güçlendirmek",
        ],
        bullet_number_id,
    )

    add_section(document, "İçerik sözlüğü")
    add_label_value(
        document,
        "Anahtar kelimeler",
        "nitelikli kahve, kahve demleme, V60, ev baristası, sürdürülebilir kahve",
    )
    add_label_value(
        document,
        "Kaçınılacak kelimeler",
        "en ucuz, mucize, kesin sonuç, dünyanın en iyisi",
    )

    add_section(document, "Rakipler")
    add_bullets(
        document,
        [
            "Örnek Roastery",
            "Mahalle Kahvecisi",
            "Global Coffee Lab",
        ],
        bullet_number_id,
    )

    add_section(document, "Yayın tercihleri")
    add_label_value(document, "Tercih edilen platformlar", "Instagram, LinkedIn, YouTube")
    add_label_value(document, "Varsayılan içerik tonu", "Samimi, öğretici ve güven veren")
    add_label_value(document, "Dil", "Türkçe")

    add_section(document, "Ek bağlam")
    context = document.add_paragraph()
    context.paragraph_format.space_after = Pt(12)
    context_run = context.add_run(
        "Lunera Coffee bütün ürün iletişiminde izlenebilirlik ve adil üretici ilişkilerini öne çıkarır. İndirim odaklı iletişim yerine ürünün hikâyesini ve kullanım deneyimini anlatır. Her içerikte agresif satış çağrısı yapmak yerine yararlı bir bilgi veya uygulanabilir küçük bir öneri sunar."
    )
    set_run_font(context_run, 10.5, BODY)

    closing = document.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.paragraph_format.space_after = Pt(0)
    shade_paragraph(closing, "F4F4F5")
    set_paragraph_border(closing, "top", LINE, size=8, space=6)
    closing_run = closing.add_run(
        "İpucu: Briefiniz ne kadar açık ve güncelse, KadeAI'nin ürettiği içerik de markanıza o kadar tutarlı yaklaşır."
    )
    set_run_font(closing_run, 10, MUTED, italic=True)

    document.core_properties.title = "KadeAI Şirket Briefi Örneği"
    document.core_properties.subject = "KadeAI şirket hafızası için örnek marka briefi"
    document.core_properties.author = "Kade New Media"
    document.core_properties.keywords = "KadeAI, şirket briefi, marka hafızası, içerik"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
